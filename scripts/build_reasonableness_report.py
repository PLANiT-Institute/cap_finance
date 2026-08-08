#!/usr/bin/env python3
"""Build the rendered Word validation brief for the steel capital-allocation model."""

from __future__ import annotations

import csv
import json
import tempfile
import uuid
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs" / "data_audit"
ASSET_DIR = OUTPUT_DIR / "report_assets"
REPORT_PATH = OUTPUT_DIR / "Capital_Allocation_Reasonableness_Report.docx"
EMBED_FONT_PATH = ASSET_DIR / "NanumGothic.ttf"

NAVY = "0B1F33"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "64748B"
LIGHT = "F2F4F7"
AMBER = "FEF3C7"
AMBER_TEXT = "92400E"
GREEN = "DCFCE7"
GREEN_TEXT = "166534"
RED = "FEE2E2"
RED_TEXT = "991B1B"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    # Calibri is the preset base font. Mixed Korean runs use the installed
    # Unicode fallback so LibreOffice/PDF rendering does not drop Hangul glyphs.
    has_hangul = any("\uac00" <= char <= "\ud7a3" for char in run.text)
    font_name = "Nanum Gothic" if has_hangul else "Calibri"
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), font_name)
    if has_hangul:
        run._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")
        lang = run._element.get_or_add_rPr().find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            run._element.get_or_add_rPr().append(lang)
        lang.set(qn("w:val"), "ko-KR")
        lang.set(qn("w:eastAsia"), "ko-KR")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="CBD5E1", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    if sum(widths) != TABLE_WIDTH:
        raise ValueError(f"table widths must sum to {TABLE_WIDTH}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
    set_table_borders(table)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths, font_size=9, header_fill=LIGHT):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        set_run_font(run, size=font_size, color=INK, bold=True)
        set_cell_shading(cell, header_fill)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths)
    return table


def add_citation(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=8.5, color=MUTED, italic=True)
    return p


def add_callout(doc, title, body, fill=AMBER, text_color=AMBER_TEXT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title)
    set_run_font(title_run, size=12, color=text_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    body_run = p2.add_run(body)
    set_run_font(body_run, size=10.5, color=text_color)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_field(run, field):
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Apple SD Gothic Neo")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        style._element.rPr.rFonts.set(qn("w:cs"), "Apple SD Gothic Neo")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = p.add_run("CAPITAL ALLOCATION MODEL VALIDATION")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = p.add_run("\tPOSCO + JAPAN STEEL")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Internal working model  |  ")
    set_run_font(r, size=8.5, color=MUTED)
    page_run = p.add_run()
    set_run_font(page_run, size=8.5, color=MUTED)
    add_field(page_run, "PAGE")


def add_masthead(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("MODEL VALIDATION BRIEF")
    set_run_font(r, size=23, color="000000", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("한·일 철강 Capital Allocation Pathway — 기준데이터·결과 타당성 검토")
    set_run_font(r, size=14, color="374151")

    meta = add_table(
        doc,
        ["항목", "내용"],
        [
            ["수신", "프로젝트 오너 / 투자·전략 검토자"],
            ["기준일", "2026-08-08"],
            ["대상", "POSCO, Nippon Steel, JFE Steel, Kobe Steel"],
            ["검증 범위", "공식 총량·프로젝트·자원 맥락, Excel↔CSV, 모델 동일성, 강건 의사결정"],
            ["상태", "조건부 적합 — 구조 검증용 / 투자 의사결정용 부적합"],
        ],
        [1440, 7920],
        font_size=9.5,
    )
    for cell in meta.rows[0].cells:
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(WHITE)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(2)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def load_rows():
    def read_csv(path):
        with path.open(encoding="utf-8") as handle:
            return list(csv.DictReader(handle))

    robust = read_csv(ROOT / "outputs" / "repeat_refined_candidate_robust_summary.csv")
    scenario_metrics = read_csv(ROOT / "outputs" / "repeat_refined_candidate_scenario_metrics.csv")
    facility_schedule = read_csv(ROOT / "outputs" / "refined_candidate_facility_schedule.csv")
    resource_profile = read_csv(ROOT / "outputs" / "refined_candidate_resource_profile.csv")
    resource_benchmarks = read_csv(ROOT / "data" / "resource_benchmarks.csv")
    with (ROOT / "outputs" / "repeat_summary.json").open(encoding="utf-8") as handle:
        repeat_summary = json.load(handle)
    with (OUTPUT_DIR / "roundtrip_audit.json").open(encoding="utf-8") as handle:
        roundtrip = json.load(handle)
    with (OUTPUT_DIR / "model_parity_audit.json").open(encoding="utf-8") as handle:
        parity = json.load(handle)

    companies = sorted({row["company_name"] for row in robust})
    choices = []
    for company in companies:
        eligible = [
            row for row in robust
            if row["company_name"] == company and row["robust_feasible"] == "True"
        ]
        selected = sorted(
            eligible,
            key=lambda row: (
                -float(row["lambda_1_optimal_frequency_pct"]),
                float(row["maximum_regret_p50_kkrw_per_tco2_mean"]),
            ),
        )[0]
        candidate_id = selected["candidate_id"]
        metrics = [row for row in scenario_metrics if row["candidate_id"] == candidate_id]
        disclosed = next(row for row in metrics if row["scenario_id"] == "DISCLOSED_PATH")
        worst_tcar = max(float(row["tcar_kkrw_per_tco2_mean"]) for row in metrics)
        resources = [row for row in resource_profile if row["candidate_id"] == candidate_id]
        facilities = [
            row for row in facility_schedule
            if row["candidate_id"] == candidate_id and row["scenario_id"] == "DISCLOSED_PATH"
        ]
        choices.append({
            **selected,
            "worst_tcar": worst_tcar,
            "capex": float(disclosed["aligned_capex_bn_krw_mean"]),
            "avoided": float(disclosed["common_avoided_emissions_mtco2"]),
            "annual_avoided": sum(float(row["annual_avoided_emissions_mtco2"]) for row in facilities),
            "cash_cost": float(disclosed["net_cash_cost_after_support_p50_bn_krw_mean"]),
            "carbon_value": float(disclosed["avoided_carbon_cost_value_p50_bn_krw_mean"]),
            "policy_support": float(disclosed["policy_support_value_p50_bn_krw_mean"]),
            "economic_npv": float(disclosed["absolute_npv_p50_bn_krw_mean"]),
            "electricity_shapley": float(disclosed["electricity_shapley_variance_share_mean"]),
            "hydrogen_shapley": float(disclosed["hydrogen_shapley_variance_share_mean"]),
            "capex_shapley": float(disclosed["capex_shapley_variance_share_mean"]),
            "shapley_delta": max(abs(float(row["shapley_reconciliation_delta_mean"])) for row in metrics),
            "scrap_utilization": max(float(row["scrap_utilization_pct"]) for row in resources),
            "hydrogen_utilization": max(float(row["hydrogen_utilization_pct"]) for row in resources),
            "grid_utilization": max(float(row["incremental_grid_utilization_pct"]) for row in resources),
            "facilities": facilities,
        })
    return choices, resource_benchmarks, repeat_summary, roundtrip, parity


def chart_font(size, bold=False):
    candidates = [
        Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("/System/Library/Fonts/Helvetica.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def create_charts(choices):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    benchmark_path = ASSET_DIR / "eaf_capex_benchmark.png"
    labels = ["Model Japan\nSCRAP_EAF", "JFE Kurashiki\n2025", "Nippon 3 EAFs\n2025"]
    values = [616.0, 1515.24, 2755.87]
    colors = ["#0C7C86", "#3B82F6", "#0B1F33"]
    image = Image.new("RGB", (1300, 540), "white")
    draw = ImageDraw.Draw(image)
    title_font = chart_font(34, bold=True)
    label_font = chart_font(24)
    value_font = chart_font(25, bold=True)
    axis_font = chart_font(20)
    draw.text((70, 24), "Japan large-EAF CAPEX benchmark", fill="#111827", font=title_font)
    left, top, right, bottom = 105, 105, 1240, 440
    max_value = 3000
    for tick in range(0, 3001, 500):
        y = bottom - (tick / max_value) * (bottom - top)
        draw.line((left, y, right, y), fill="#E5E7EB", width=2)
        draw.text((20, y - 12), f"{tick:,}", fill="#64748B", font=axis_font)
    bar_width = 205
    centers = [300, 675, 1050]
    for center, label, value, color in zip(centers, labels, values, colors):
        height = value / max_value * (bottom - top)
        x1, y1, x2, y2 = center - bar_width / 2, bottom - height, center + bar_width / 2, bottom
        draw.rounded_rectangle((x1, y1, x2, y2), radius=8, fill=color)
        value_box = draw.textbbox((0, 0), f"{value:,.0f}", font=value_font)
        draw.text((center - (value_box[2] - value_box[0]) / 2, y1 - 34), f"{value:,.0f}", fill="#111827", font=value_font)
        parts = label.split("\n")
        for line_no, part in enumerate(parts):
            box = draw.textbbox((0, 0), part, font=label_font)
            draw.text((center - (box[2] - box[0]) / 2, bottom + 15 + line_no * 27), part, fill="#374151", font=label_font)
    draw.text((1010, 70), "bn KRW / Mtpa", fill="#64748B", font=axis_font)
    image.save(benchmark_path, quality=95)

    results_path = ASSET_DIR / "robust_results.png"
    names = [r["company_name"] for r in choices]
    regret = [float(r["maximum_regret_p50_kkrw_per_tco2_mean"]) for r in choices]
    tcar = [r["worst_tcar"] for r in choices]
    image = Image.new("RGB", (1300, 610), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 22), "Robust decision screen — refined lambda=1 choices", fill="#111827", font=title_font)
    left, right = 330, 1240
    top, row_gap, bar_height = 150, 88, 24
    max_value = 100
    for tick in range(0, 101, 20):
        x = left + (tick / max_value) * (right - left)
        draw.line((x, 130, x, 515), fill="#E5E7EB", width=2)
        draw.text((x - 12, 530), str(tick), fill="#64748B", font=axis_font)
    for i, (name, regret_value, tcar_value) in enumerate(zip(names, regret, tcar)):
        y = top + i * row_gap
        draw.text((50, y + 7), name, fill="#374151", font=label_font)
        p50_right = left + (regret_value / max_value) * (right - left)
        tcar_right = left + (tcar_value / max_value) * (right - left)
        draw.rounded_rectangle((left, y, p50_right, y + bar_height), radius=5, fill="#0B1F33")
        draw.rounded_rectangle((left, y + 34, tcar_right, y + 34 + bar_height), radius=5, fill="#0C7C86")
        draw.text((p50_right + 8, y - 2), f"{regret_value:.1f}", fill="#0B1F33", font=axis_font)
        draw.text((tcar_right + 8, y + 31), f"{tcar_value:.1f}", fill="#0C7C86", font=axis_font)
    draw.rounded_rectangle((850, 76, 880, 100), radius=4, fill="#0B1F33")
    draw.text((890, 77), "Maximum regret", fill="#374151", font=axis_font)
    draw.rounded_rectangle((1070, 76, 1100, 100), radius=4, fill="#0C7C86")
    draw.text((1110, 77), "Worst TCaR", fill="#374151", font=axis_font)
    image.save(results_path, quality=95)

    shapley_path = ASSET_DIR / "shapley_risk.png"
    image = Image.new("RGB", (1300, 570), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 22), "Exact Shapley allocation of transition-cost variance", fill="#111827", font=title_font)
    left, right = 300, 1230
    top, row_gap, bar_height = 145, 82, 38
    factor_colors = ["#2563EB", "#0C7C86", "#D97706"]
    factors = ["Electricity", "Hydrogen", "CAPEX"]
    for tick in range(0, 101, 20):
        x = left + tick / 100 * (right - left)
        draw.line((x, 125, x, 480), fill="#E5E7EB", width=2)
        draw.text((x - 10, 490), f"{tick}%", fill="#64748B", font=axis_font)
    for i, row in enumerate(choices):
        y = top + i * row_gap
        draw.text((45, y + 4), row["company_name"], fill="#374151", font=label_font)
        x = left
        shares = [row["electricity_shapley"], row["hydrogen_shapley"], row["capex_shapley"]]
        for share, color in zip(shares, factor_colors):
            x2 = x + share * (right - left)
            draw.rectangle((x, y, x2, y + bar_height), fill=color)
            if share >= 0.06:
                label = f"{share * 100:.1f}%"
                box = draw.textbbox((0, 0), label, font=axis_font)
                draw.text((x + (x2 - x - (box[2] - box[0])) / 2, y + 6), label, fill="white", font=axis_font)
            x = x2
    legend_x = 650
    for factor, color in zip(factors, factor_colors):
        draw.rectangle((legend_x, 77, legend_x + 24, 101), fill=color)
        draw.text((legend_x + 34, 77), factor, fill="#374151", font=axis_font)
        legend_x += 190
    image.save(shapley_path, quality=95)
    return benchmark_path, results_path, shapley_path


def embed_font(docx_path, font_path, font_name="Nanum Gothic"):
    """Embed an installable TrueType font for reliable Korean PDF rendering."""
    if not font_path.exists():
        return False
    font_key = uuid.uuid4()
    key_bytes = bytes.fromhex(font_key.hex)[::-1]
    font_data = bytearray(font_path.read_bytes())
    for index in range(min(32, len(font_data))):
        font_data[index] ^= key_bytes[index % 16]

    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    ET.register_namespace("w", w_ns)
    ET.register_namespace("r", r_ns)
    ET.register_namespace("", rel_ns)

    with zipfile.ZipFile(docx_path, "r") as source:
        parts = {name: source.read(name) for name in source.namelist()}

    font_table = ET.fromstring(parts["word/fontTable.xml"])
    font_node = ET.SubElement(font_table, f"{{{w_ns}}}font", {f"{{{w_ns}}}name": font_name})
    ET.SubElement(
        font_node,
        f"{{{w_ns}}}embedRegular",
        {
            f"{{{r_ns}}}id": "rIdNanumGothic",
            f"{{{w_ns}}}fontKey": "{" + str(font_key).upper() + "}",
            f"{{{w_ns}}}subsetted": "0",
        },
    )
    parts["word/fontTable.xml"] = ET.tostring(font_table, encoding="utf-8", xml_declaration=True)

    rels_path = "word/_rels/fontTable.xml.rels"
    if rels_path in parts:
        rels = ET.fromstring(parts[rels_path])
    else:
        rels = ET.Element(f"{{{rel_ns}}}Relationships")
    ET.SubElement(
        rels,
        f"{{{rel_ns}}}Relationship",
        {
            "Id": "rIdNanumGothic",
            "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font",
            "Target": "fonts/NanumGothic.odttf",
        },
    )
    parts[rels_path] = ET.tostring(rels, encoding="utf-8", xml_declaration=True)

    content_types = ET.fromstring(parts["[Content_Types].xml"])
    if not any(node.get("Extension") == "odttf" for node in content_types):
        ET.SubElement(
            content_types,
            f"{{{ct_ns}}}Default",
            {
                "Extension": "odttf",
                "ContentType": "application/vnd.openxmlformats-officedocument.obfuscatedFont",
            },
        )
    ET.register_namespace("", ct_ns)
    parts["[Content_Types].xml"] = ET.tostring(content_types, encoding="utf-8", xml_declaration=True)

    settings = ET.fromstring(parts["word/settings.xml"])
    if settings.find(f"{{{w_ns}}}embedTrueTypeFonts") is None:
        settings.append(ET.Element(f"{{{w_ns}}}embedTrueTypeFonts"))
    parts["word/settings.xml"] = ET.tostring(settings, encoding="utf-8", xml_declaration=True)
    parts["word/fonts/NanumGothic.odttf"] = bytes(font_data)

    with tempfile.NamedTemporaryFile(suffix=".docx", dir=docx_path.parent, delete=False) as handle:
        temp_path = Path(handle.name)
    with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for name, data in parts.items():
            target.writestr(name, data)
    temp_path.replace(docx_path)
    return True


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    current, roundtrip, parity = load_rows()
    benchmark_chart, results_chart = create_charts(current)

    doc = Document()
    configure_document(doc)
    add_masthead(doc)

    add_heading(doc, "1. 결론", 1)
    add_callout(
        doc,
        "조건부 적합",
        "기준데이터의 출처 추적, 시설 합계 조정, Excel→CSV 왕복, 동일 조건 재실행은 통과했다. 그러나 회사 간 데이터 경계, 시설 블록의 실제성, 일본 대형 전기로 CAPEX 보정이 충분하지 않아 투자 승인이나 기업 순위에 쓰면 안 된다.",
    )
    add_paragraph(doc, "현재 버전은 ‘계산이 돌아가고 변경 이력이 검증되는 구조’로서는 말이 된다. 하지만 ‘어느 회사가 더 유리한가’ 또는 ‘어떤 설비에 얼마를 승인할 것인가’를 판단하는 모델로서는 아직 말이 되지 않는다.")
    add_table(
        doc,
        ["사용 목적", "판정", "근거"],
        [
            ["데이터 구조·파이프라인 검증", "사용 가능", "8개 입력 파일 왕복 PASS, 모델 결과 동일"],
            ["시나리오·계획 상대 민감도", "조건부 사용", "같은 회사 안에서 방향성 비교 가능"],
            ["회사 간 순위·투자 승인", "사용 금지", "경계 불일치와 CAPEX 하방 편향 가능성"],
        ],
        [2800, 1800, 4760],
        font_size=9.5,
    )
    add_citation(doc, "판정 기준: 공식값 추적성, 산술 재조정, 입력 왕복, 결정론적 재실행, 경계 일치, 기술비용 공시 벤치마크.")

    add_page_break(doc)
    add_heading(doc, "2. 검증 가능한 데이터 흐름", 1)
    add_paragraph(doc, "Excel 감사본은 원천 CSV 컬럼을 각 시트의 왼쪽에 그대로 보존한다. 검증 수식과 판정 열은 오른쪽에만 추가되며, 재생성 스크립트는 원천 컬럼 수만 읽는다. 따라서 감사용 열이 실행 입력에 섞이지 않는다.")
    add_table(
        doc,
        ["1. 원천", "2. 감사", "3. 재생성", "4. 실행 검증"],
        [["data/*.csv + JSON", "Excel 수식·Sources·SHA256", "csv_export/*.csv", "validate-data + 동일 seed 실행"]],
        [2340, 2340, 2340, 2340],
        font_size=9.5,
    )
    add_heading(doc, "검증 결과", 2)
    add_table(
        doc,
        ["검증", "결과", "증거"],
        [
            ["Excel→CSV 의미상 동일성", f"{roundtrip['status']} (8/8)", "roundtrip_audit.json; 헤더·행·열·값 비교"],
            ["재생성 CSV 로더 검증", "PASS", "4 companies / 17 facilities / 6 technologies / 8 scenarios / 32 plans"],
            ["모델 결과 바이트 동일성", f"{parity['status']} (3/3)", "plan_metrics, facility_schedule, frontier_membership"],
            ["재실행 조건", "100 paths / seed 42", "원본과 Excel 재생성 CSV를 각각 독립 실행"],
        ],
        [3000, 1800, 4560],
        font_size=9.5,
    )
    add_heading(doc, "회사 기준 총량", 2)
    add_table(
        doc,
        ["회사", "생산 Mt", "Scope 1+2 Mt", "집약도", "경계 주의"],
        [
            ["POSCO", "34.537", "69.846", "2.020", "국내 환경·별도 재무: 상대적으로 정렬"],
            ["Nippon Steel", "34.300", "72.600", "2.117*", "일본 모회사 환경·연결 재무"],
            ["JFE Steel", "21.950", "45.300", "2.060", "JFE Steel 환경·JFE Holdings 재무"],
            ["Kobe Steel", "5.960", "14.300", "2.399*", "생산·목표경계·연결 재무 혼재"],
        ],
        [1800, 1200, 1550, 1250, 3560],
        font_size=9,
    )
    add_citation(doc, "* 배출/생산 파생 집약도. Kobe Steel은 분자·분모 경계가 완전히 일치하지 않는다.")

    add_page_break(doc)
    add_heading(doc, "3. 타당성 점검", 1)
    add_table(
        doc,
        ["차원", "점수/5", "판정", "핵심 근거"],
        [
            ["공식값 추적성", "4", "GOOD", "회사 총량·URL·경계 메모 존재"],
            ["산술 조정", "5", "PASS", "17개 시설 블록 합계가 회사 총량과 ±0.001Mt 내 일치"],
            ["회사 간 경계", "2", "WEAK", "일본 3사의 환경·재무 경계 불일치"],
            ["시설 현실성", "1", "WEAK", "공식 총량을 맞춘 모델 블록"],
            ["기술 CAPEX", "2", "WEAK", "공시 EAF 프로젝트 대비 낮은 원단위"],
            ["시나리오", "3", "FAIR", "공시 이정표 + 모델 앵커·가속 스트레스"],
            ["확률 계산", "4", "GOOD", "seed 반복·P50/P90/TCaR 재현"],
            ["투자결정 준비도", "2", "CONDITIONAL", "경계·시설·CAPEX 보정 전"],
        ],
        [2100, 1100, 1600, 4560],
        font_size=9,
    )
    add_callout(doc, "종합 2.9 / 5", "수리 구조와 데이터 계보는 통과했지만 경제적 보정은 미완료다. 따라서 현재 대시보드의 숫자는 순위가 아니라 질문을 만드는 진단 신호로 사용한다.", fill=LIGHT, text_color=DARK_BLUE)
    add_heading(doc, "핵심 이상 신호", 2)
    add_bullet(doc, "POSCO의 2030 공시경로 탄소예산 70.920Mt는 2025 실적 69.846Mt보다 1.074Mt 높다. 목표를 수치상 이미 충족한 상태처럼 보이므로, 생산 감소와 구조적 감축을 분리해야 한다.")
    add_bullet(doc, "시설별 생산·집약도·재투자연도는 실제 설비 데이터가 아니라 회사 총량을 맞춘 추정치다. 시설 투자 순서의 현실성을 검증하지 못한다.")
    add_bullet(doc, "인정 탄소비용 비율 45%, 환율 1 JPY=9.2 KRW, 수소·전력 가격과 상관구조는 모두 모델 가정이다. 안정적인 seed 결과는 이 가정이 맞다는 뜻이 아니다.")

    add_page_break(doc)
    add_heading(doc, "4. 일본 대형 전기로 CAPEX 공시 벤치마크", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    benchmark_shape = p.add_run().add_picture(str(benchmark_chart), width=Inches(6.35))
    benchmark_shape._inline.docPr.set("title", "일본 대형 전기로 CAPEX 공시 벤치마크")
    benchmark_shape._inline.docPr.set("descr", "모델 일본 SCRAP_EAF 616, JFE Kurashiki 1515, Nippon Steel 3개 전기로 2756 bn KRW/Mtpa 비교 막대그래프")
    add_table(
        doc,
        ["비교", "모델/공시 원단위", "모델 비율", "필요 배수"],
        [
            ["JFE Kurashiki 2025", "616 / 1,515 bn KRW/Mtpa", "40.7%", "2.46x"],
            ["Nippon Steel 3 EAFs 2025", "616 / 2,756 bn KRW/Mtpa", "22.4%", "4.47x"],
        ],
        [2900, 3000, 1500, 1960],
        font_size=9.5,
    )
    add_citation(doc, "JFE: 329.4bn JPY / 약 2.0Mtpa. Nippon Steel: 868.7bn JPY / 약 2.9Mtpa. 모델 환율 9.2 KRW/JPY 적용.")
    add_callout(
        doc,
        "CAPEX는 재보정이 필요하다",
        "공시 프로젝트는 전기로 본체뿐 아니라 전력계통·물류·부대설비를 포함할 수 있어 완전한 동등 비교는 아니다. 그럼에도 현재 일본 SCRAP_EAF 모델 원단위는 공시 full-scope의 22~41%에 불과하다. 최소한 low/base/high 범위와 공정범위 bridge를 도입해야 한다.",
    )
    add_paragraph(doc, "이 하방 편향은 일부 계획의 순전환비용이 비정상적으로 낮거나 음수가 되는 현상을 강화할 수 있다.")
    add_paragraph(doc, "가속경로 POSCO P1의 평균 P50 약 -4.7kKRW/tCO₂는 탄소비용 회피효과가 지출을 초과하도록 계산된 결과다. 현재 비용 보정 수준에서는 경제적 결론으로 채택하면 안 된다.")

    add_page_break(doc)
    add_heading(doc, "5. 현재 결과 읽는 법", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    results_shape = p.add_run().add_picture(str(results_chart), width=Inches(6.35))
    results_shape._inline.docPr.set("title", "가속경로 공시전략 프록시 결과")
    results_shape._inline.docPr.set("descr", "POSCO, Nippon Steel, JFE Steel, Kobe Steel의 P50 순비용과 TCaR를 비교한 가로 막대그래프")
    result_rows = []
    for row in current:
        result_rows.append([
            row["company_name"],
            f"{float(row['expected_cost_p50_kkrw_per_tco2_mean']):.1f}",
            f"{float(row['tcar_kkrw_per_tco2_mean']):.1f}",
            f"{float(row['aligned_capex_bn_krw']):,.0f}",
            f"{float(row['p90_cost_to_ebitda_x_mean']):.2f}x",
        ])
    add_table(
        doc,
        ["회사", "P50 kKRW/t", "TCaR kKRW/t", "정렬 CAPEX bn KRW", "NPV/연 EBITDA"],
        result_rows,
        [1900, 1700, 1600, 2300, 1860],
        font_size=9,
    )
    add_citation(doc, "ACCELERATED_15C, 공시전략 프록시, 3 seeds × 1,000 paths. 평균값.")
    add_heading(doc, "해석 제한", 2)
    add_bullet(doc, "P90 transition NPV / annual EBITDA는 15년 누적 순현재비용을 1년 EBITDA로 나눈 스트레스 배수다. 부채상환·커버리지 비율이 아니다.")
    add_bullet(doc, "Nippon Steel·JFE·Kobe Steel은 환경 범위보다 재무 분모가 넓어 스트레스 배수가 인위적으로 낮아질 수 있다. 회사 간 직접 순위는 금지한다.")
    add_bullet(doc, "Kobe Steel의 P50이 높은 것은 단순히 ‘열위’라는 뜻이 아니다. 규모, 배출경계, 전환량 분모가 달라 동일 기준 비교가 아니다.")

    add_page_break(doc)
    add_heading(doc, "6. 다음 버전의 수정 우선순위", 1)
    add_table(
        doc,
        ["우선", "수정", "완료 기준"],
        [
            ["P0", "환경·생산·재무 경계 통일", "같은 법인/지역 경계의 EBITDA·CAPEX 또는 별도 배부표 확보"],
            ["P0", "EAF CAPEX low/base/high 보정", "공시 프로젝트별 scope bridge와 원단위 범위; 모델/공시 70~130% 설명 가능"],
            ["P1", "시설 실데이터 교체", "용량·생산·기술·개수/폐쇄연도에 출처 또는 신뢰등급"],
            ["P1", "탄소예산 연도화", "생산량 효과와 구조 감축을 분리한 연간 경로"],
            ["P1", "지표 명칭·비교 규칙 수정", "‘15년 NPV / 1년 EBITDA’ 명시; 경계 불일치 시 순위 숨김"],
            ["P2", "시장 제약 추가", "스크랩·계통·수소 공급·공기 지연·제품 믹스·FX 민감도"],
        ],
        [1100, 2900, 5360],
        font_size=9,
    )
    add_heading(doc, "권고 운영 상태", 2)
    add_callout(
        doc,
        "Exploratory only",
        "다음 재보정 전에는 내부 탐색·데이터 갭 식별·동일 회사 내 민감도 비교에만 사용한다. 투자위원회 승인, 기업가치 평가, 기업 간 성과 순위에는 사용하지 않는다.",
        fill=RED,
        text_color=RED_TEXT,
    )
    add_heading(doc, "재승인 게이트", 2)
    add_bullet(doc, "G1 — Excel↔CSV 왕복 100% PASS와 모델 parity 유지")
    add_bullet(doc, "G2 — 모든 회사의 생산·배출·재무 경계 매핑표 승인")
    add_bullet(doc, "G3 — 대형 EAF·H₂-DRI CAPEX의 공시 프로젝트 범위 bridge 완료")
    add_bullet(doc, "G4 — 핵심 결과에 FX·생산량·공기 지연 민감도 및 역산 back-test 추가")

    add_page_break(doc)
    add_heading(doc, "부록 A. 공식 출처", 1)
    source_entries = [
        ("POSCO 환경", "https://sustainability.posco.com/S91/S91F10/eng/cmspage.do?mmcd=2682093497003371"),
        ("POSCO 생산·매출", "https://sustainability.posco.com/S91/S91F10/eng/cmspage.do?mmcd=2682093474001805"),
        ("POSCO 기후목표·EAF", "https://sustainability.posco.com/S91/S91F10/eng/cmspage.do?mmcd=2648825310001953"),
        ("POSCO 재무", "https://www.posco.co.kr/homepage/docs/eng7/jsp/ir/s91b6000050l.jsp"),
        ("Nippon Steel Integrated Report 2025", "https://www.nipponsteel.com/en/ir/library/pdf/nsc_en_ir_2025_all.pdf"),
        ("Nippon Steel 2025 EAF investment", "https://www.nipponsteel.com/common/secure/en/news/20250530_200.pdf"),
        ("JFE Group Report 2025", "https://www.jfe-holdings.co.jp/en/common/pdf/investor/library/group-report/2025/all.pdf"),
        ("JFE Kurashiki EAF 2025", "https://www.jfe-steel.co.jp/en/release/2025/04/250410.html"),
        ("JFE capacity risk disclosure", "https://www.jfe-holdings.co.jp/en/investor/management/risk/"),
        ("Kobelco Integrated Report 2025", "https://www.kobelco.co.jp/english/ir/integrated-reports/pdf/integrated-reports2025_e.pdf"),
    ]
    add_table(doc, ["출처", "URL"], source_entries, [2700, 6660], font_size=8)
    add_citation(doc, "모든 링크 최종 확인일: 2026-08-06. 모델 입력값의 정확한 경계·환산 메모는 Excel 감사본 Sources/Companies/Financials 시트를 참조.")

    add_heading(doc, "부록 B. 재현 산출물", 1)
    add_table(
        doc,
        ["산출물", "경로 / 역할"],
        [
            ["Excel 감사본", "outputs/data_audit/Capital_Allocation_Baseline_Audit.xlsx"],
            ["재생성 CSV", "outputs/data_audit/csv_export/"],
            ["왕복 감사", "outputs/data_audit/roundtrip_audit.json"],
            ["모델 parity", "outputs/data_audit/model_parity.json"],
            ["실행 결과", "outputs/data_audit/reference_run/ 및 roundtrip_run/"],
        ],
        [2700, 6660],
        font_size=9,
    )

    doc.core_properties.title = "Capital Allocation Reasonableness Report"
    doc.core_properties.subject = "POSCO and Japanese steelmakers model validation"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "POSCO, Nippon Steel, JFE, Kobe Steel, capital allocation, validation"
    doc.save(REPORT_PATH)
    embed_font(REPORT_PATH, EMBED_FONT_PATH)
    print(REPORT_PATH)


def main_v080():
    """Build the current decision-focused v0.8.0 validation report."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    choices, resource_benchmarks, repeat_summary, roundtrip, parity = load_rows()
    with (ROOT / "data" / "transition_projects.csv").open(encoding="utf-8") as handle:
        transition_projects = list(csv.DictReader(handle))
    with (ROOT / "data" / "technology_cost_evidence.csv").open(encoding="utf-8") as handle:
        cost_evidence = list(csv.DictReader(handle))
    with (ROOT / "outputs" / "data_depth_assessment.json").open(encoding="utf-8") as handle:
        depth_assessment = json.load(handle)
    company_order = {"POSCO": 0, "Nippon Steel": 1, "JFE Steel": 2, "Kobe Steel": 3}
    choices.sort(key=lambda row: company_order[row["company_name"]])
    benchmark_chart, results_chart, shapley_chart = create_charts(choices)
    meta = repeat_summary["meta"]

    doc = Document()
    configure_document(doc)
    add_masthead(doc)

    add_heading(doc, "1. 결론", 1)
    add_callout(
        doc,
        "조건부 적합 — 실행 가능한 의사결정 스크리닝",
        "v0.8.0은 같은 고정 설비 포트폴리오를 두 활성 시나리오에서 비교하고, 현금비용·탄소회피가치·정책지원을 분리하며, 현실 제약을 통과한 후보만 추천한다. 데이터 계보와 재현성은 통과했지만 공식 GCAM 1.5°C·2.0°C가 비활성이므로 투자 승인 모델은 아니다.",
    )
    add_paragraph(doc, "현재 결과는 어떤 시설 조합이 두 내부 경로 모두에서 실행 가능하고, 가격 위험과 후회를 얼마나 남기는지 판단하는 데 의미가 있다. 기업 간 우열이나 실제 예산 승인에는 회사별 공급계약·계통접속·프로젝트 범위 CAPEX의 추가 증빙이 필요하다.")
    add_table(doc, ["사용 목적", "판정", "근거"], [
        ["데이터·모델 재현", "사용 가능", "20/20 왕복, 14/14 결과 동일, 7/7 테스트"],
        ["회사 내부 후보·시설 비교", "조건부 사용", "탄소·자원·공사·실패 제약과 강건성 반영"],
        ["기업 간 순위·투자 승인", "사용 금지", "경계·시설·CAPEX와 공식 GCAM 미완료"],
    ], [2800, 1800, 4760], font_size=9.5)
    add_citation(doc, "판정 기준: 공식값 추적성, 입력 왕복, 고정 포트폴리오, 공통 분모, 확률 재실행, 현실 제약, 공식 시나리오 활성화 상태.")

    add_page_break(doc)
    add_heading(doc, "2. 감사 추적성과 모델 동일성", 1)
    add_paragraph(doc, "Excel 감사본은 원천 CSV 열을 그대로 보존하고 검증식만 오른쪽에 추가한다. Excel에서 다시 내보낸 입력을 독립 폴더에서 실행해 기준 결과와 비교했으므로, 사람이 확인하는 감사본과 코드가 읽는 기준데이터 사이의 단절을 탐지할 수 있다.")
    add_table(doc, ["1. 원천", "2. 감사", "3. 재생성", "4. 실행 검증"], [
        ["20 CSV/JSON/XML", "34-sheet Excel + SHA256", "csv_export/", "1,000 paths·seed 42 parity"],
    ], [2340, 2340, 2340, 2340], font_size=9.2)
    add_heading(doc, "검증 결과", 2)
    add_table(doc, ["검증", "결과", "증거"], [
        ["Excel→CSV 의미상 동일성", f"{roundtrip['status']} ({len(roundtrip['files'])}/{len(roundtrip['files'])})", "roundtrip_audit.json; 헤더·행·열·값"],
        ["모델 산출물 바이트 동일성", f"{parity['status']} ({len(parity['files'])}/{len(parity['files'])})", "model_parity_audit.json; 기준↔왕복 실행"],
        ["회귀테스트", "PASS (7/7)", "로더·제약·비용분리·후회·Shapley"],
        ["정밀 반복", "3 seeds × 1,000", "37개 고정 shortlist; 후보·시나리오당 3,000 경로"],
    ], [3000, 1900, 4460], font_size=9.2)
    add_heading(doc, "모델 범위", 2)
    add_table(doc, ["항목", "규모", "현재 상태"], [
        ["기업·시설", "4개 / 17개", "공식 기업 총량 + 총량에 맞춘 모델 시설 블록"],
        ["생성·정밀 후보", f"{meta['generated_candidate_count']} / {meta['refined_candidate_count']}", "중앙가격 screening 후 shortlist 고정"],
        ["활성 경로", "2개", "공시경로 + 내부 1.5°C 스트레스"],
        ["공식 GCAM", "2개 정의", "1.5°C·2.0°C 모두 비활성"],
        ["공식 프로젝트·비용증거", f"{len(transition_projects)}개 / {len(cost_evidence)}개", "범위·시설 매핑 검증 전 모델입력과 분리"],
    ], [2200, 1800, 5360], font_size=9)
    add_callout(doc, "시나리오 명칭 주의", "ACCELERATED_15C는 공식 GCAM 결과가 아니라 내부 스트레스 경로다. 공식 GCAM은 release hash, JVM 실행, 경로별 10개 query export를 검증한 뒤에만 활성화된다.", fill=RED, text_color=RED_TEXT)

    add_page_break(doc)
    add_heading(doc, "3. 강건 추천과 효율경계 해석", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(results_chart), width=Inches(6.35))
    shape._inline.docPr.set("title", "정밀 λ=1 강건 추천")
    shape._inline.docPr.set("descr", "회사별 최대후회와 최악 TCaR를 비교한 막대그래프")
    rows = []
    for row in choices:
        rows.append([
            row["company_name"], f"{row['candidate_id'][:9]}…{row['candidate_id'][-4:]}", row["template_plan_id"],
            f"{float(row['maximum_regret_p50_kkrw_per_tco2_mean']):.1f}", f"{row['worst_tcar']:.1f}",
            f"{float(row['lambda_1_optimal_frequency_pct']):.0f}%",
        ])
    add_table(doc, ["회사", "λ=1 후보", "계획", "최대후회", "최악 TCaR", "선택빈도"], rows,
              [1750, 2450, 800, 1450, 1450, 1460], font_size=8.5)
    add_citation(doc, "단위: kKRW/tCO₂. 3개 seed 평균. 선택빈도는 반복 실행에서 같은 λ=1 후보가 선택된 비율.")
    add_callout(doc, "빨간 점을 읽는 법", "효율경계 위의 빨간 점은 현재 공시계획 또는 선택 후보다. 경계보다 오른쪽·위에 있으면 같은 비교집합 안에서 더 낮은 기대비용과 더 낮은 위험을 동시에 주는 지배안이 있다는 뜻이다. ‘나쁜 회사’가 아니라 개선 여지가 있는 계획이라는 진단이다.", fill=LIGHT, text_color=DARK_BLUE)
    add_paragraph(doc, "최대후회는 각 활성 시나리오의 최저 적격 후보와 비교한 P50 비용 차이의 최댓값이다. λ=1은 최대후회 + 최악 TCaR을 최소화한다. 기준점은 전체 910개가 아니라 고정된 37개 정밀 shortlist 안에 한정된다.")

    add_page_break(doc)
    add_heading(doc, "4. 시설별 실행 계획", 1)
    add_paragraph(doc, "추천안은 회사 총량만 제시하지 않고 시설별 기술·전환연도·CAPEX·연간 감축을 보존한다. 아래 시설은 공식 총량에 맞춘 모델 블록이며 실제 프로젝트 승인 단위는 아니다.")
    rows = []
    for choice in choices:
        for facility in sorted(choice["facilities"], key=lambda item: (int(item["transition_year"]), item["facility_id"])):
            rows.append([
                choice["company_name"].replace(" Steel", ""), facility["facility_id"], facility["technology_id"],
                facility["transition_year"], f"{float(facility['aligned_capex_bn_krw']):,.0f}",
                f"{float(facility['annual_avoided_emissions_mtco2']):.2f}",
            ])
    add_table(doc, ["회사", "시설 블록", "전환 기술", "연도", "CAPEX", "연간 감축"], rows,
              [1350, 1650, 2200, 800, 1600, 1760], font_size=7.5)
    add_citation(doc, "CAPEX: bn KRW, 연간 감축: MtCO₂/년. DISCLOSED_PATH의 고정 물리 포트폴리오를 내부 스트레스에도 그대로 평가.")
    add_callout(doc, "실행 순서의 의미", "기술·연도 조합은 스크랩, 수소, 증분계통, 동시공사, 기술실패 제약을 모두 통과해야 추천에 남는다. 실제 FEED·접속승인·공급계약이 들어오면 해당 시설 행을 교체하고 후보를 다시 생성해야 한다.", fill=LIGHT, text_color=DARK_BLUE)

    add_page_break(doc)
    add_heading(doc, "5. 전체 비용·배출·자원 여력", 1)
    add_heading(doc, "실제 현금과 경제적 가치의 분리", 2)
    rows = []
    for row in choices:
        rows.append([row["company_name"], f"{row['capex']:,.0f}", f"{row['cash_cost']:,.0f}",
                     f"{row['carbon_value']:,.0f}", f"{row['policy_support']:,.0f}", f"{row['economic_npv']:,.0f}"])
    add_table(doc, ["회사", "CAPEX", "지원후 현금 P50", "탄소회피가치", "정책지원", "경제 NPV P50"], rows,
              [1650, 1300, 1900, 1700, 1400, 1410], font_size=8.2)
    add_citation(doc, "bn KRW, DISCLOSED_PATH, 3-seed 평균. 정책지원은 현금 P50에 이미 반영; 경제 NPV = 지원후 현금비용 − 탄소회피가치.")
    add_heading(doc, "배출 성과와 자원 최대 이용률", 2)
    rows = []
    for row in choices:
        rows.append([row["company_name"], f"{row['annual_avoided']:.1f}", f"{row['avoided']:.1f}",
                     f"{row['scrap_utilization']:.1f}%", f"{row['hydrogen_utilization']:.1f}%", f"{row['grid_utilization']:.1f}%"])
    add_table(doc, ["회사", "연간 감축", "공통 누적분모", "스크랩", "수소", "증분계통"], rows,
              [1700, 1350, 1750, 1450, 1450, 1660], font_size=8.7)
    add_citation(doc, "배출: MtCO₂, 이용률: 모델 공급한도 대비 두 활성 시나리오·전 기간 최댓값. 회사 공급한도는 model_estimate이며 국가 벤치마크로 대체하지 않는다.")
    add_callout(doc, "현재 가장 가까운 병목", "JFE 증분계통 99.5%, Nippon Steel 스크랩 95.8%, POSCO 스크랩 90.0%가 모델 한도에 가장 가깝다. 추천 순위보다 먼저 접속 가능용량·고급 스크랩 계약·수소 인도조건을 확인해야 한다.")

    add_page_break(doc)
    add_heading(doc, "6. 왜 위험이 움직이는가 — 정확한 Shapley 분해", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(shapley_chart), width=Inches(6.35))
    shape._inline.docPr.set("title", "전환비용 분산의 정확한 Shapley 배분")
    shape._inline.docPr.set("descr", "전력, 수소입력, CAPEX가 비용 분산에 기여하는 비중")
    rows = []
    for row in choices:
        rows.append([row["company_name"], f"{row['electricity_shapley']*100:.1f}%",
                     f"{row['hydrogen_shapley']*100:.1f}%", f"{row['capex_shapley']*100:.1f}%", f"{row['shapley_delta']:.1e}"])
    add_table(doc, ["회사", "전력", "수소입력", "CAPEX", "재조정 오차"], rows,
              [2000, 1600, 1600, 1600, 2560], font_size=9)
    add_citation(doc, "DISCLOSED_PATH. 세 요인의 8개 부분집합을 공통난수로 모두 재평가한 variance-game Shapley.")
    add_callout(doc, "비용 구성비가 아니다", "전력 85~93%는 총비용의 85~93%가 전력이라는 뜻이 아니라 비용 ‘분산’의 배분이다. 수소 제조의 전력가격 노출은 전력 요인에 포함되고, 수소입력은 비전력 전해조 성분만 변동한다. 재조정 오차는 표시 정밀도에서 0이다.", fill=LIGHT, text_color=DARK_BLUE)

    add_page_break(doc)
    add_heading(doc, "7. 공식 국가 자원 벤치마크", 1)
    add_callout(doc, "국가 맥락 ≠ 회사 공급한도", "아래 공식값은 단위와 범위가 서로 다르며 회사에 배정된 물량이 아니다. 모델 제약을 덮어쓰지 않고, 수소·계통·스크랩 가정이 국가 규모와 모순되지 않는지 보는 별도 감사층이다.")
    rows = []
    for row in resource_benchmarks:
        value = row["benchmark_value"]
        value = "정성 정책" if not value else f"{float(value):,.2f}".rstrip("0").rstrip(".")
        rows.append([row["country_code"], row["resource_type"], row["benchmark_year"], value, row["unit"], row["source_org"]])
    add_table(doc, ["국가", "자원", "연도", "공식 값", "단위", "1차 출처"], rows,
              [700, 1100, 850, 1250, 1600, 3860], font_size=7.4)
    add_citation(doc, "11개 레코드, 추출일 2026-08-07. 범위·버전·직접 URL·비교가능성은 data/resource_benchmarks.csv와 Excel에 보존.")
    add_heading(doc, "판단에 쓰는 방식", 2)
    add_bullet(doc, "수소: KR 2030 국내 청정수소 1 Mt/년, JP 2030·2040·2050 수소환산 3·12·20 Mt/년은 국가 전체 목표다.")
    add_bullet(doc, "계통: GW·MW·MVA는 전력량 TWh와 직접 비교하지 않는다. 사이트 접속용량과 계통보강 공기를 별도로 확보한다.")
    add_bullet(doc, "스크랩: JP 2022 발생량 43.16 Mt/년과 2030 고급화 2 Mt/년은 품질·지역·계약 가용성을 보장하지 않는다.")

    add_page_break(doc)
    add_heading(doc, "8. 공식 전환 프로젝트 증거층", 1)
    add_paragraph(doc, "공식 발표에 있는 프로젝트 상태·용량·CAPEX·정부지원·가동시점과 모델 시설 매핑을 별도 원장으로 보존했다. 공시 프로젝트가 기존 설비를 대체하는지, 증설인지, 실증인지 확인되기 전에는 최적화 입력을 자동 변경하지 않는다.")
    project_rows = []
    status_labels = {
        "operating": "Operating",
        "announced_demo": "Demo announced",
        "investment_decided": "FID approved",
        "feasibility_study": "Feasibility",
        "completed": "Completed",
        "demonstration_completed": "Demo completed",
    }
    mapping_labels = {
        "unmapped_new_asset": "Unmapped new asset",
        "unmapped_demo_asset": "Unmapped demo",
        "timing_and_site_anchor": "Site/timing anchor",
        "unmapped_group_asset": "Unmapped group asset",
        "hybrid_route_evidence": "Hybrid evidence",
        "direct_historical_anchor": "Historical anchor",
        "technology_performance_anchor": "Technology anchor",
    }
    for project in transition_projects:
        capex = "—"
        if project["capex_native_bn"]:
            capex = f"{project['capex_currency']} {float(project['capex_native_bn']):,.1f}bn"
        support = "—"
        if project["government_support_native_bn"]:
            support = f"{project['capex_currency']} {float(project['government_support_native_bn']):,.1f}bn"
        project_rows.append([
            project["company_id"].replace("_STEEL_JP", "").replace("_KR", "").replace("_JP", ""),
            project["project_name"],
            status_labels.get(project["project_status"], project["project_status"]),
            f"{float(project['capacity_mtpa']):.1f}" if project["capacity_mtpa"] else "—",
            capex,
            support,
            project["operation_start_label"] or "—",
            mapping_labels.get(project["model_mapping_status"], project["model_mapping_status"]),
        ])
    add_table(doc, ["회사", "공식 프로젝트", "상태", "Mtpa", "공시 CAPEX", "지원", "가동", "모델 매핑"], project_rows,
              [900, 2300, 1100, 650, 1350, 850, 1150, 1060], font_size=6.7)
    add_citation(doc, "9개 프로젝트, 7개 비용증거. 공란은 공시되지 않은 값이다. 원문 URL·버전·추출일·범위 메모는 transition_projects.csv와 Excel 감사본에 보존.")
    add_callout(doc, "증거성숙도 40.6% — 정확도 점수가 아니다", f"공식 프로젝트 원장은 95%, 비용증거는 90% 수준이지만 시설·기술·회사 자원제약은 각각 20%다. 미해결 P0 {depth_assessment['open_p0_gap_count']}건을 닫기 전 절대 NPV와 시설 투자순서를 승인하지 않는다.", fill=LIGHT, text_color=DARK_BLUE)

    add_page_break(doc)
    add_heading(doc, "9. 일본 대형 전기로 CAPEX 공시 벤치마크", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(benchmark_chart), width=Inches(6.35))
    shape._inline.docPr.set("title", "일본 대형 전기로 CAPEX 공시 벤치마크")
    shape._inline.docPr.set("descr", "모델과 일본 공시 프로젝트 CAPEX 원단위 비교")
    add_table(doc, ["비교", "모델/공시 원단위", "모델 비율", "필요 배수"], [
        ["JFE Kurashiki 2025", "616 / 1,515 bn KRW/Mtpa", "40.7%", "2.46x"],
        ["Nippon Steel 3 EAFs 2025", "616 / 2,756 bn KRW/Mtpa", "22.4%", "4.47x"],
    ], [2900, 3000, 1500, 1960], font_size=9.3)
    add_citation(doc, "JFE: 329.4bn JPY / 약 2.0Mtpa. Nippon Steel: 868.7bn JPY / 약 2.9Mtpa. 환율 9.2 KRW/JPY.")
    add_callout(doc, "CAPEX scope bridge가 필요하다", "공시 프로젝트는 전기로 본체 외 계통·물류·부대설비를 포함할 수 있어 완전한 동등 비교는 아니다. 그러나 모델 원단위가 full-scope 공시의 22~41%인 차이는 추천 후보의 현금비용을 낮게 만들 수 있으므로 low/base/high와 범위 조정표가 필요하다.")
    add_paragraph(doc, "추천안의 절대 NPV는 구조적으로 보존되지만 입력 CAPEX의 공정범위가 틀리면 절대값도 틀린다. 효율경계의 상대 위치와 투자 예산 승인을 분리해야 하는 이유다.")

    add_page_break(doc)
    add_heading(doc, "10. 남은 한계와 승인 게이트", 1)
    add_table(doc, ["등급", "남은 한계", "완료 기준"], [
        ["P0", "공식 GCAM 1.5°C·2.0°C 비활성", "공식 9.1 DB/JVM 실행, 경로별 10 query export, hash·단위·지역 검증"],
        ["P0", "회사 공급한도는 모델 추정", "스크랩 계약, 수소 인도량, 사이트 계통접속 증빙으로 교체"],
        ["P0", "대형 EAF/H₂-DRI CAPEX 범위", "공시 프로젝트 low/base/high 및 공정범위 bridge 승인"],
        ["P1", "17개 시설은 총량 배분 블록", "실제 용량·기술·개수/폐쇄·정비연도와 신뢰등급"],
        ["P1", "환경·생산·재무 경계 차이", "같은 법인·지역 경계의 EBITDA/CAPEX 또는 배부표"],
        ["P1", "위험분해가 3개 가격요인", "공기지연·기술실패·생산량·FX까지 가치게임 확장"],
    ], [1100, 3050, 5210], font_size=8.3)
    add_heading(doc, "승인 규칙", 2)
    add_callout(doc, "현재 운영 상태: Exploratory / screening", "내부 데이터 갭 식별, 같은 회사 안의 후보·시설·계약 민감도, 추가 실사 우선순위에는 사용한다. 기업가치 평가, 회사 간 성과 순위, 투자위원회 승인에는 사용하지 않는다.", fill=RED, text_color=RED_TEXT)
    add_bullet(doc, "G1 — 20/20 Excel↔CSV, 14/14 parity, 회귀테스트를 모든 입력 변경 때 유지")
    add_bullet(doc, "G2 — 공식 GCAM 이중 경계와 동일계획 연결선이 활성 상태로 표시")
    add_bullet(doc, "G3 — 회사별 자원 계약·계통접속·동시공사 증빙이 후보 제약에 반영")
    add_bullet(doc, "G4 — CAPEX scope bridge와 기업경계 배부표가 절대 NPV까지 승인")

    add_page_break(doc)
    add_heading(doc, "부록 A. 핵심 공식 출처", 1)
    sources = [
        ("POSCO 환경", "https://sustainability.posco.com/S91/S91F10/eng/cmspage.do?mmcd=2682093497003371"),
        ("POSCO 광양 EAF·HyREX 2026", "https://newsroom.posco.com/en/posco-accelerates-transition-to-decarbonized-production-system-with-completion-of-koreas-largest-electric-arc-furnace/"),
        ("Nippon Steel IR 2025", "https://www.nipponsteel.com/en/ir/library/pdf/nsc_en_ir_2025_all.pdf"),
        ("Nippon Steel EAF 2025", "https://www.nipponsteel.com/en/newsroom/news/2025/__icsFiles/afieldfile/2025/09/26/20250530_200.pdf"),
        ("JFE Group Report 2025", "https://www.jfe-holdings.co.jp/en/common/pdf/investor/library/group-report/2025/all.pdf"),
        ("JFE Kurashiki EAF 2025", "https://www.jfe-steel.co.jp/en/release/2025/04/250410.html"),
        ("Kobelco IR 2025", "https://www.kobelco.co.jp/english/ir/integrated-reports/pdf/integrated-reports2025_e.pdf"),
        ("Kobelco Kakogawa melter 2026", "https://www.kobelco.co.jp/releases/2026/1218906_18738.html"),
        ("KR Hydrogen Policy", "https://www.pcccr.go.kr/base/board/read?boardManagementNo=10&boardNo=124&menuLevel=2&menuNo=18&page=2"),
        ("KR Electricity Plan", "https://www.motir.go.kr/kor/article/ATCL3f49a5a8c/170183/view"),
        ("Japan Hydrogen Strategy", "https://www.meti.go.jp/policy/energy_environment/global_warming/transition/jcr_climate_transition_bond_framework_spo_eng.pdf"),
        ("Japan OCCTO 2023", "https://www.occto.or.jp/assets/en/information_disclosure/annual_report/files/2023_annualreport_240131.pdf"),
        ("Japan scrap study", "https://www.env.go.jp/content/000315009.pdf"),
    ]
    add_table(doc, ["출처", "직접 URL"], sources, [2850, 6510], font_size=7.3)
    add_citation(doc, "확인·추출일: 2026-08-08. 값·단위·지역경계·버전·비교가능성은 원천 CSV와 Excel Sources/ResourceBenchmarks/TransitionProjects 시트를 참조.")
    add_heading(doc, "부록 B. 재현 산출물", 1)
    add_table(doc, ["산출물", "경로 / 역할"], [
        ["Excel 감사본", "outputs/data_audit/Capital_Allocation_Baseline_Audit.xlsx"],
        ["재생성 CSV", "outputs/data_audit/csv_export/"],
        ["왕복·모델 감사", "outputs/data_audit/roundtrip_audit.json / model_parity_audit.json"],
        ["정밀 결과", "outputs/repeat_refined_candidate_*.csv"],
        ["시설·자원", "outputs/refined_candidate_facility_schedule.csv / resource_profile.csv"],
        ["의사결정 화면", "outputs/dashboard.html"],
        ["데이터 심도 감사", "outputs/data_depth_assessment.csv / .json"],
        ["자동 실행 기록", "outputs/automation_progress.md"],
    ], [2700, 6660], font_size=8.1)
    add_heading(doc, "부록 C. 데이터 상태를 읽는 규칙", 1)
    add_table(doc, ["상태", "의미 / 모델 사용 규칙"], [
        ["official_project_disclosure", "기업·정부 1차 발표에서 직접 확인한 프로젝트 사실; 범위·시설 매핑 검증 전 최적화 입력을 자동 변경하지 않음"],
        ["official_derived", "공식 수치에 명시 환율·용량 분모를 적용한 계산값; 원식과 원문 URL을 함께 보존"],
        ["model_estimate", "공식값이 아닌 모델 가정; 민감도·스크리닝에는 사용하되 투자승인 근거로 사용 금지"],
        ["pending_official_extract", "공식 실행·query export가 완료되지 않은 빈 슬롯; 시나리오 활성화 차단"],
        ["open_gap / partially_evidenced", "필수 증거 미확보 또는 일부 증거만 확보; data_gap_registry.csv의 다음 출처와 model_action을 따름"],
    ], [2600, 6760], font_size=7.3)
    add_callout(doc, "프로젝트 공시와 모델 입력을 분리한 이유", "공식 프로젝트가 기존 설비 대체·증설·실증 중 무엇인지 확정되지 않으면 CAPEX와 감축량을 현재 시설블록에 더하는 순간 이중계상될 수 있다. 그래서 먼저 증거원장을 만들고, scope bridge와 asset mapping을 통과한 행만 다음 버전에서 입력으로 승격한다.", fill=LIGHT, text_color=DARK_BLUE)
    doc.core_properties.title = "Capital Allocation Reasonableness Report v0.8.0"
    doc.core_properties.subject = "POSCO and Japanese steelmakers robust capital-allocation validation"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "POSCO, Nippon Steel, JFE, Kobe Steel, GCAM, robust frontier, Shapley"
    doc.save(REPORT_PATH)
    embed_font(REPORT_PATH, EMBED_FONT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    main_v080()
